import json
import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Cm
from docx.text.run import Font
def check_docx_format(file_path):
    # 打开文档
    document = Document(file_path)
    # 定义要检查的样式
    styles_to_check = {
        1: {'name': 'Heading 1', 'font': {'name': '黑体', 'size': Pt(16), 'bold': True}, 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
        2: {'name': 'Heading 2', 'font': {'name': '黑体', 'size': Pt(12), 'bold': True}, 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
        3: {'name': 'Heading 3', 'font': {'name': '宋体', 'size': Pt(12), 'bold': True}, 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
        'list': {'name': 'List Paragraph', 'font': {'name': '宋体', 'size': Pt(12)}, 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
        'normal': {'name': 'Normal', 'font': {'name': '宋体', 'size': Pt(12)}, 'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY, 'indent': Cm(0.74)},
        'header': {'name': '页眉', 'font': {'name': '宋体', 'size': Pt(10)}, 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
        'footer': {'name': '页脚', 'font': {'name': '宋体', 'size': Pt(10)}, 'alignment': WD_ALIGN_PARAGRAPH.CENTER}
    }
    errors = []
    
    defaultFontName = "default Font"

    #两个层次
    # 一开始是para.style
    # 然后是run styles
    for para in document.paragraphs:
        StyleFontSize = para.style.font.size
        StyleFontName = para.style.font.name
        if para.style.name == styles_to_check[1]['name']:
            #对runs进行检查
            for run in para.runs:
                fontSize = StyleFontSize if StyleFontSize!=None else run.style.font.size
                fontName = StyleFontName if StyleFontName!=None else run.font.name
                if fontName == None:
                    fontName = defaultFontName
                if fontName != styles_to_check[1]['font']['name'] or \
                        fontSize != styles_to_check[1]['font']['size'] or \
                        para.style.font.bold != styles_to_check[1]['font']['bold'] or \
                        para.style.paragraph_format.alignment != styles_to_check[1]['alignment']:
                    errors.append({'heading_level': '一级标题', 'font': fontName, 'size': fontSize.pt if fontSize!=None else "None", 'alignment': para.style.paragraph_format.alignment, "text":run.text})
        elif para.style.name == styles_to_check[2]['name']:
            for run in para.runs:
                fontSize = StyleFontSize if StyleFontSize!=None else run.style.font.size
                fontName = StyleFontName if StyleFontName!=None else run.font.name
                if fontName == None:
                    fontName = defaultFontName
                if fontName != styles_to_check[2]['font']['name'] or \
                        fontSize != styles_to_check[2]['font']['size'] or \
                        para.style.font.bold != styles_to_check[2]['font']['bold'] or \
                        para.style.paragraph_format.alignment != styles_to_check[2]['alignment']:
                    errors.append({'heading_level': '二级标题','font': fontName, 'size': fontSize.pt if fontSize!=None else "None", 'alignment': para.style.paragraph_format.alignment, "text":run.text})

        elif para.style.name == styles_to_check[3]['name']:
            for run in para.runs:
                fontSize = StyleFontSize if StyleFontSize!=None else run.style.font.size
                fontName = StyleFontName if StyleFontName!=None else run.font.name
                if fontName == None:
                    fontName = defaultFontName
                if fontName != styles_to_check[3]['font']['name'] or \
                        fontSize != styles_to_check[3]['font']['size'] or \
                        para.style.font.bold != styles_to_check[3]['font']['bold'] or \
                        para.style.paragraph_format.alignment != styles_to_check[3]['alignment']:
                    errors.append({'heading_level': '三级标题', 'font': fontName, 'size': fontSize.pt if fontSize!=None else "None", 'alignment': para.style.paragraph_format.alignment,"text":run.text})

        elif para.style.name == styles_to_check['list']['name']:
            for run in para.runs:
                fontSize = StyleFontSize if StyleFontSize!=None else run.style.font.size
                fontName = StyleFontName if StyleFontName!=None else run.font.name
                if fontName == None:
                    fontName = defaultFontName
                if fontName != styles_to_check['list']['font']['name'] or \
                    fontSize != styles_to_check['list']['font']['size'] or \
                    para.style.paragraph_format.alignment != styles_to_check['list']['alignment']:
                    errors.append({'heading_level': '条、款、项标题', 'font': fontName, 'size': fontSize.pt if fontSize!=None else "None", 'alignment': para.style.paragraph_format.alignment,"text":run.text})
        # 对正文进行检测
        elif para.style.name == styles_to_check['normal']['name']:
            for run in para.runs:
                # 宋体 小四号 两端对齐 首行缩进 2 字符
                fontSize = StyleFontSize if StyleFontSize!=None else run.style.font.size
                fontName = StyleFontName if StyleFontName!=None else run.font.name
                if fontName == None:
                    fontName = defaultFontName
                if fontName != styles_to_check['normal']['font']['name'] or  \
                    fontSize != styles_to_check['normal']['font']['size'] or \
                    para.style.paragraph_format.alignment != styles_to_check['normal']['alignment'] or \
                    para.style.paragraph_format.first_line_indent != styles_to_check['normal']['indent']:
                    errors.append({'heading_level': '正文', 'font': fontName, 'size': fontSize.pt if fontSize!=None else "None", 'alignment': para.style.paragraph_format.alignment,"text":run.text})

    # 检查页眉和页脚的格式是否符合要求
    # 检查每个节的页眉和页脚
    for section in document.sections:
        # 检查页眉格式
        if section.header is not None:
            for paragraph in section.header.paragraphs:
                # 检查字体格式
                for run in paragraph.runs:
                    if run.bold or run.underline:
                        errors.append({'message':'错误：页眉中存在不符合格式要求的字体格式'})
                # 检查段落对齐方式
                if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    errors.append({'message':'错误：页眉中存在不符合格式要求的段落对齐方式'})
                # 检查字体大小
                for run in paragraph.runs:
                    if run.font.size != Pt(5):
                        errors.append({'message':'错误：页眉中存在不符合格式要求的字体大小'})
                        
        # 检查页脚格式
        if section.footer is not None:
            for paragraph in section.footer.paragraphs:
                # 检查字体格式
                for run in paragraph.runs:
                    if run.bold or run.underline:
                        errors.append({'message':'错误：页脚中存在不符合格式要求的字体格式'})
                # 检查段落对齐方式
                if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    errors.append({'message':'错误：页脚中存在不符合格式要求的段落对齐方式'})
                # 检查字体大小
                for run in paragraph.runs:
                    if run.font.size != Pt(5):
                        errors.append({'message':'错误：页脚中存在不符合格式要求的字体大小'})

    return errors


def getFormat(path):
    res = []
    errors = check_docx_format(path)
    for error in errors:
        res.append({
            "message": json.dumps(error),
            "paragraph": error.get("text", "")
        })
    return res

